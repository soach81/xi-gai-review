"""
Import additional questions from 题库背诵版 and 知识点及练习题 files.
Merges into existing questions.json, deduping by question text similarity.
"""
import re
import json
from pathlib import Path
from docx import Document

DATA_DIR = Path(r"C:\Users\99114\Desktop\习概复习\data")
QUESTIONS_PATH = DATA_DIR / "questions.json"

# Map 讲 names to chapter IDs
JIANG_TO_CHAPTER = {
    "第一讲": "guide",
    "第二讲": "ch01",
    "第三讲": "ch02",
    "第四讲": "ch03",
    "第五讲": "ch04",
    "第六讲": "ch05",
    "第七讲": "ch06",
    "第八讲": "ch07",
    "第九讲": "ch08",
    "第十讲": "ch09",
    "第十一讲": "ch10",
    "第十二讲": "ch11",
    "第十三讲": "ch12",
    "第十四讲": "ch13",
    "第十五讲": "ch14",
    "第十六讲": "ch15",
}

LECTURE_NUM_MAP = {
    "第一讲": "guide", "第二讲": "ch01", "第三讲": "ch02",
    "第四讲": "ch03", "第五讲": "ch04", "第六讲": "ch05",
    "第七讲": "ch06", "第八讲": "ch07", "第九讲": "ch08",
    "第十讲": "ch09", "第十一讲": "ch10", "第十二讲": "ch11",
    "第十三讲": "ch12", "第十四讲": "ch13", "第十五讲": "ch14",
    "第十六讲": "ch15",
}


def detect_lecture(text):
    for jname in LECTURE_NUM_MAP:
        if text.startswith(jname):
            return jname
    return None


def parse_answer_line(text):
    """Parse answer line like '参考答案： C' or '1.D  2.D  3.D'"""
    # Single answer format: "参考答案： C" or similar
    match = re.search(r'参考\s*答案\s*[：:]\s*([A-G]+)', text)
    if match:
        return [l for l in match.group(1)]
    # Multi-answer format: "1.D   2.D   3.D"
    answers = []
    for m in re.finditer(r'\d+\.([A-G]+)', text):
        answers.extend(list(m.group(1)))
    return answers


def parse_baosong(doc):
    """Parse 题库背诵版 docx"""
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    questions = []
    current_chapter_id = None
    current_qtype = None
    current_qtext = None
    current_options = []
    current_answers = []
    i = 0

    while i < len(paragraphs):
        text = paragraphs[i]

        # Detect lecture
        jname = detect_lecture(text)
        if jname:
            current_chapter_id = LECTURE_NUM_MAP[jname]
            current_qtype = None
            i += 1
            continue

        # Detect question type
        if "单选题" in text and ("一、" in text[:10] or "、" in text[:5]):
            current_qtype = "single"
            i += 1
            continue
        if "多选题" in text and ("二、" in text[:10] or "、" in text[:5]):
            current_qtype = "multi"
            i += 1
            continue

        # Detect answer key - skip
        if "参考答案" in text or re.match(r'^\d+[\.\、][A-G]', text):
            i += 1
            continue
        # Skip section headers
        if re.match(r'^[三四五六]、', text) or "简答题" in text or "判断题" in text or "论述题" in text:
            current_qtype = None
            i += 1
            continue

        # Parse question
        if current_qtype and current_chapter_id:
            q_match = re.match(r'^(\d+)[\.\、\s\)、]+(.+)', text)
            if q_match:
                q_text = q_match.group(2).strip()
                options = []
                j = i + 1
                while j < len(paragraphs):
                    next_text = paragraphs[j]
                    opt_match = re.match(r'^([A-G])\s*[\.\、\）\)]\s*(.+)', next_text)
                    if opt_match:
                        options.append(opt_match.group(2).strip())
                        j += 1
                        continue
                    # Check if next is a new question, answer key, or section header
                    if re.match(r'^\d+[\.\、\s\)]', next_text):
                        break
                    if "参考答案" in next_text or re.match(r'^[A-G][\s\.\、]', next_text):
                        j += 1
                        break
                    if re.match(r'^[三四五六]、', next_text) or "单选题" in next_text or "多选题" in next_text:
                        break
                    # Continuation of question text
                    q_text += " " + next_text
                    j += 1

                if options and current_chapter_id:
                    # Find answer in the answer key (nearby in the doc)
                    # For 背诵版, answers are embedded inline as single-letter after question
                    # We'll parse them from the answer key section
                    questions.append({
                        "chapter_id": current_chapter_id,
                        "type": current_qtype,
                        "question": q_text[:500],
                        "options": options,
                        "source": "题库背诵版"
                    })
                i = j
                continue
        i += 1

    return questions


def parse_answer_file(baosong_qs, answer_docx_path):
    """Try to match answers from an answer key file."""
    # For now, answers are embedded in the 背诵版 as the last option-like line
    # Or in separate answer sections. We'll need to match manually.
    return baosong_qs


def main():
    # Load existing
    existing = json.loads(QUESTIONS_PATH.read_text(encoding='utf-8'))
    existing_texts = set()
    for ch in existing['chapters']:
        for q in ch['questions']:
            existing_texts.add(q['question'][:60])

    # Parse 题库背诵版
    baosong_path = Path(r"C:\Users\99114\Desktop\习概\02_题库_主要\题库背诵版.docx")
    doc = Document(baosong_path)
    new_qs = parse_baosong(doc)
    print(f"Parsed {len(new_qs)} questions from 题库背诵版")

    # Parse 知识点及练习题 (16 files)
    jiang_dir = Path(r"C:\Users\99114\Desktop\习概\03_题库_辅助\习思想概论知识点及练习题（1-16讲全）")
    for fpath in sorted(jiang_dir.glob("*")):
        if fpath.suffix not in ('.docx', '.doc'):
            continue
        try:
            doc = Document(fpath)
            qs = parse_baosong(doc)  # Same format
            new_qs.extend(qs)
        except Exception as e:
            print(f"  Skip {fpath.name}: {e}")

    print(f"Total new questions (before dedup): {len(new_qs)}")

    # Deduplicate and assign IDs
    merged_count = 0
    qid_counter = sum(len(ch['questions']) for ch in existing['chapters'])
    chapter_map = {ch['id']: ch for ch in existing['chapters']}

    for nq in new_qs:
        ch_id = nq.get('chapter_id', 'guide')
        if ch_id not in chapter_map:
            continue
        # Check dedup by first 60 chars of question
        if nq['question'][:60] in existing_texts:
            continue
        existing_texts.add(nq['question'][:60])
        qid_counter += 1
        chapter_map[ch_id]['questions'].append({
            "id": f"{ch_id}_q{qid_counter:04d}",
            "type": nq['type'],
            "question": nq['question'][:500],
            "options": nq['options'],
            "answer": [],  # Answers will need matching from answer keys
            "source": nq.get('source', '辅助题库'),
            "priority": 2
        })
        merged_count += 1

    # Update total
    existing['meta']['totalQuestions'] = sum(len(ch['questions']) for ch in existing['chapters'])
    existing['meta']['version'] = '1.1'

    QUESTIONS_PATH.write_text(json.dumps(existing, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"Merged {merged_count} new questions")
    print(f"Total: {existing['meta']['totalQuestions']} questions")

    for ch in existing['chapters']:
        print(f"  {ch['id']}: {ch['name']} - {len(ch['questions'])} questions")


if __name__ == "__main__":
    main()
