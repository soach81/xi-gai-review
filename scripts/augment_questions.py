"""
Augment questions.json with auxiliary question banks.
Sources:
  1. 16讲 docx files (1-11 available as docx, 12-16 are .doc → skip)
  2. 习近平新时代中国特色社会主义思想概论题库(1).docx (Moodle quiz format)
  3. 习概23题库.pdf.docx (with inline 【答案】)

Deduplication: difflib.SequenceMatcher > 85% similarity → duplicate.
Priority: keep "23版题库" source; otherwise keep first encountered.
"""
import re
import json
import difflib
from pathlib import Path
from docx import Document

# ─── Paths ───────────────────────────────────────────────────
EXISTING_JSON = Path(r"C:\Users\99114\Desktop\习概复习\data\questions.json")
SRC_16LECTURES = Path(r"C:\Users\99114\Desktop\习概\03_题库_辅助\习思想概论知识点及练习题（1-16讲全）")
SRC_TIKU1 = Path(r"C:\Users\99114\Desktop\习概\03_题库_辅助\习近平新时代中国特色社会主义思想概论题库(1).docx")
SRC_23TIKU = Path(r"C:\Users\99114\Desktop\习概\03_题库_辅助\习概23题库.pdf.docx")

# ─── Chapter mapping ──────────────────────────────────────────
# Lecture title → chapter_id
LECTURE_TO_CHAPTER = {
    "马克思主义中国化新的飞跃": "guide",
    "坚持和发展中国特色社会主义的总任务": "ch01",
    "坚持党的全面领导": "ch03",
    "坚持以人民为中心": "ch04",
    "以新发展理念引领高质量发展": "ch06",
    "全面深化改革": "ch05",
    "发展全过程人民民主": "ch08",
    "全面依法治国": "ch09",
    "建设社会主义文化强国": "ch10",
    "加强以民生为重点的社会建设": "ch11",
    "建设社会主义生态文明": "ch12",
}

# 23题库 讲 → chapter_id mapping
JIANG_TO_CHAPTER = {
    "第一讲": "guide",
    "第二讲": "ch01",
    "第三讲": "ch03",
    "第四讲": "ch04",
    "第五讲": "ch06",
    "第六讲": "ch05",
    "第七讲": "ch08",
    "第八讲": "ch09",
    "第九讲": "ch10",
    "第十讲": "ch11",
    "第十一讲": "ch12",
    "第十二讲": "ch13",
    "第十三讲": "ch14",
    "第十四讲": "ch15",
    "第十五讲": "ch16",
    "第十六讲": "ch17",
}

# Chapter names for matching headers in 23题库
JIANG_NAMES = list(JIANG_TO_CHAPTER.keys())

# ─── Helpers ──────────────────────────────────────────────────

def option_to_index(letter):
    """A→0, B→1, etc."""
    return ord(letter.upper()) - ord('A')


def parse_answer_letters(text):
    """Extract answer letters from text like '正确答案：D' or '【答案】ABC' or '1.D 2.D 3.D' or 'A,B,C,D'"""
    # Try comma-separated: A,B,C,D
    if ',' in text and re.search(r'[A-G]', text):
        parts = text.replace('，', ',').split(',')
        letters = [p.strip() for p in parts if p.strip() and re.match(r'^[A-G]$', p.strip())]
        if letters:
            return letters

    # Try space-separated with number prefix: 1.D 2.D 3.D
    numbered = re.findall(r'\d+\s*[\.\、]?\s*([A-G])', text)
    if numbered:
        return numbered

    # Try plain letters: ABC or A B C
    plain = re.findall(r'[A-G]', text)
    # If it's an answer key line like "1.D   2.D   3.D" extract the letters
    if plain and '正确答案' in text:
        return plain

    return plain if plain else []


def similarity(a, b):
    """Calculate similarity between two question texts."""
    return difflib.SequenceMatcher(None, a, b).ratio()


def normalize_question(text):
    """Normalize question text for comparison."""
    # Remove all whitespace
    return re.sub(r'\s+', '', text)


# ─── Parse existing questions.json ────────────────────────────

def load_existing():
    with open(EXISTING_JSON, 'r', encoding='utf-8') as f:
        return json.load(f)


# ─── Parse 16-lecture files (docx format) ─────────────────────

def parse_16lecture_file(filepath):
    """Parse a single 16-lecture docx and return list of (chapter_id, question_dict)."""
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"  SKIP: Cannot read {filepath.name}: {e}")
        return None, []

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Detect lecture title
    lecture_title = None
    for text in paragraphs:
        m = re.match(r'第[一二三四五六七八九十]+讲\s*(.+)', text)
        if m:
            lecture_title = m.group(1).strip()
            break
        # Also try format without number: just the title
        for key in LECTURE_TO_CHAPTER:
            if key in text and len(text) < 80:
                lecture_title = key
                break
        if lecture_title:
            break

    if not lecture_title:
        print(f"  WARN: Cannot detect lecture title in {filepath.name}")
        return None, []

    # Find matching chapter
    chapter_id = None
    for key, ch_id in LECTURE_TO_CHAPTER.items():
        if key in lecture_title or lecture_title in key or lecture_title.startswith(key[:4]):
            chapter_id = ch_id
            break

    if not chapter_id:
        # Try fuzzy match
        for key, ch_id in LECTURE_TO_CHAPTER.items():
            if similarity(key, lecture_title) > 0.5:
                chapter_id = ch_id
                break

    if not chapter_id:
        print(f"  WARN: Cannot map lecture '{lecture_title}' to chapter")
        return None, []

    print(f"  Lecture: {lecture_title} → {chapter_id}")

    questions = []
    current_qtype = None  # 'single' or 'multi'

    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]

        # Detect question type section
        if '（一）单选题' in text or ('单选题' in text and ('一、' in text[:10] or '（一）' in text[:10])):
            current_qtype = "single"
            i += 1
            continue
        if '（二）多选题' in text or ('多选题' in text and ('二、' in text[:10] or '（二）' in text[:10])):
            current_qtype = "multi"
            i += 1
            continue

        # Skip non-choice sections
        if re.match(r'^(二、|[三四五六七八九十]、|（三）|（四）)', text) and ('简答' in text or '判断' in text or '论述' in text):
            current_qtype = None
            i += 1
            continue

        # Skip 参考答案 when it's a standalone answer key
        if '参考答案' in text and current_qtype:
            # Parse all answers from following text
            answer_key = text
            # Check if answers are on same line
            answer_entries = re.findall(r'(\d+)\s*[\.\、]?\s*([A-G]+)', answer_key)
            if not answer_entries:
                # Answers might be on the next line
                if i + 1 < len(paragraphs) and not re.match(r'^(一、|二、|（一）|（二）)', paragraphs[i + 1]):
                    answer_key += ' ' + paragraphs[i + 1]
                    i += 1
                answer_entries = re.findall(r'(\d+)\s*[\.\、]?\s*([A-G]+)', answer_key)

            # Backfill answers to previously parsed questions
            if answer_entries:
                for q_num_str, ans_letters in answer_entries:
                    q_num = int(q_num_str)
                    # Find the matching question (last questions without answer)
                    # Questions are numbered 1-based within each section
                    for q in reversed(questions):
                        if q.get('_qnum') == q_num and not q.get('answer'):
                            q['answer'] = [option_to_index(a) for a in ans_letters]
                            break

            i += 1
            continue

        # Parse numbered question
        if current_qtype:
            q_match = re.match(r'^(\d+)[\.\、\s\)]+(.+)', text)
            if q_match:
                q_num = int(q_match.group(1))
                q_text = q_match.group(2).strip()

                # Look ahead for options
                options = []
                j = i + 1
                while j < len(paragraphs):
                    next_text = paragraphs[j]
                    opt_match = re.match(r'^([A-G])[\.\、\）\)]\s*(.+)', next_text)
                    if opt_match:
                        options.append(opt_match.group(2).strip())
                        j += 1
                        continue
                    # Check if this is an answer or next question or section break
                    if '参考答案' in next_text:
                        break
                    if re.match(r'^\d+[\.\、\s\)]', next_text) and '参考答案' not in next_text:
                        break
                    if re.match(r'^(一、|二、|（一）|（二）|（三）|（四）)', next_text):
                        break
                    # Could be continuation of question text
                    if not re.match(r'^[A-G][\.\、\）\)]', next_text):
                        q_text += " " + next_text
                        j += 1
                    else:
                        break

                if options:
                    questions.append({
                        "_qnum": q_num,
                        "type": current_qtype,
                        "question": q_text[:500],
                        "options": options,
                        "answer": [],  # will be filled by backfill or may remain empty
                        "source": "16讲题库",
                        "priority": 2,
                    })

                i = j
                continue

        i += 1

    # Remove _qnum temporary field and filter out questions without answers
    result = []
    for q in questions:
        if q.get("answer"):
            q.pop("_qnum", None)
            result.append((chapter_id, q))

    return chapter_id, result


def parse_all_16lectures():
    """Parse all 16-lecture docx files."""
    all_questions = []
    skipped = []

    for filepath in sorted(SRC_16LECTURES.glob("*")):
        if filepath.suffix.lower() == '.doc':
            skipped.append(filepath.name)
            print(f"  SKIP .doc: {filepath.name}")
            continue
        if filepath.suffix.lower() != '.docx':
            continue

        print(f"\nParsing: {filepath.name}")
        chapter_id, questions = parse_16lecture_file(filepath)
        if questions:
            print(f"  Extracted {len(questions)} questions")
            all_questions.extend(questions)

    return all_questions, skipped


# ─── Parse 题库(1).docx (Moodle quiz format) ──────────────────

def parse_tiku1():
    """Parse 习近平新时代中国特色社会主义思想概论题库(1).docx using block-based approach."""
    try:
        doc = Document(SRC_TIKU1)
    except Exception as e:
        print(f"ERROR reading {SRC_TIKU1.name}: {e}")
        return []

    paragraphs = [p.text.strip() for p in doc.paragraphs]

    # Join all paragraphs and split by "题目N" markers
    full_text = '\n'.join(paragraphs)
    blocks = re.split(r'\n题目(\d+)\s*(?:题干)?\s*', full_text)

    questions = []
    # blocks[0] = text before first 题目, blocks[1] = number, blocks[2] = content, blocks[3] = number, etc.
    for idx in range(1, len(blocks) - 1, 2):
        q_num = blocks[idx]
        block_content = blocks[idx + 1]
        lines = [l.strip() for l in block_content.split('\n') if l.strip()]

        # Extract question text (lines before first option)
        q_parts = []
        opt_start = 0
        is_multi = False
        for j, line in enumerate(lines):
            # Check for option start
            opt_match = re.match(r'^([A-G])[\.\、\）\)]\s*(.+)', line)
            if opt_match:
                opt_start = j
                break
            # Check for 选择一项
            if '选择一项或多项' in line:
                is_multi = True
                opt_start = j + 1
                break
            if '选择一项' in line:
                opt_start = j + 1
                break
            # Not option, not selection hint → question text
            q_parts.append(line)

        q_text = ' '.join(q_parts).strip()
        if not q_text:
            continue

        # Extract options
        options_dict = {}
        for j in range(opt_start, len(lines)):
            line = lines[j]
            if '正确答案' in line:
                break
            if '选择一项' in line:
                if '多项' in line:
                    is_multi = True
                continue
            opt_match = re.match(r'^([A-G])[\.\、\）\)]\s*(.+)', line)
            if opt_match:
                letter = opt_match.group(1)
                content = opt_match.group(2).strip()
                # Remove trailing 选择一项 if attached
                content = re.sub(r'\s*选择一项.*$', '', content).strip()
                if letter not in options_dict and content:
                    options_dict[letter] = content
                continue
            # Standalone letter (Moodle artifact)
            if re.match(r'^[A-G]\s*$', line):
                continue

        # Find answer
        answer_letters = []
        for j in range(opt_start, len(lines)):
            line = lines[j]
            if '正确答案' in line:
                # Check if answer is inline: "正确答案是：A"
                ans_inline = re.search(r'正确答案[是：:]\s*(.+)', line)
                if ans_inline:
                    answer_letters = parse_answer_letters(ans_inline.group(1))
                # Otherwise check next few lines
                if not answer_letters:
                    for offset in range(1, 4):
                        if j + offset < len(lines):
                            ans_line = lines[j + offset].strip()
                            if ans_line:
                                answer_letters = parse_answer_letters(ans_line)
                                if answer_letters:
                                    break
                break

        if options_dict and answer_letters:
            options = []
            for letter in sorted(options_dict.keys()):
                options.append(options_dict[letter])

            answer_indices = [option_to_index(a) for a in answer_letters if a in options_dict]

            if answer_indices and options:
                q_type = "multi" if (is_multi or len(answer_indices) > 1) else "single"
                questions.append({
                    "type": q_type,
                    "question": q_text[:500],
                    "options": options,
                    "answer": answer_indices,
                    "source": "补充题库",
                    "priority": 2,
                })

    print(f"\nParsing 题库(1).docx: extracted {len(questions)} questions")
    return questions


# ─── Parse 习概23题库.pdf.docx ────────────────────────────────

def parse_23tiku_pdf():
    """Parse 习概23题库.pdf.docx - inline answer format."""
    try:
        doc = Document(SRC_23TIKU)
    except Exception as e:
        print(f"ERROR reading {SRC_23TIKU.name}: {e}")
        return []

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    all_questions = []
    current_chapter = None
    current_qtype = None  # 'single' or 'multi'

    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]

        # Detect chapter: 第X讲
        for jiang_name in JIANG_NAMES:
            if text.startswith(jiang_name):
                current_chapter = JIANG_TO_CHAPTER.get(jiang_name)
                print(f"  Chapter: {jiang_name} → {current_chapter}")
                i += 1
                break
        else:
            i += 1
            continue

        if not current_chapter:
            continue

        # Now parse questions within this chapter
        while i < len(paragraphs):
            text = paragraphs[i]

            # Detect next chapter
            is_next_chapter = False
            for jiang_name in JIANG_NAMES:
                if text.startswith(jiang_name) and JIANG_TO_CHAPTER.get(jiang_name) != current_chapter:
                    is_next_chapter = True
                    current_chapter = JIANG_TO_CHAPTER.get(jiang_name)
                    break
            if is_next_chapter:
                break

            # Detect question type section
            if '一、单选题' in text or '一 、单选题' in text:
                current_qtype = "single"
                i += 1
                continue
            if '二、多选题' in text or '二 、多选题' in text:
                current_qtype = "multi"
                i += 1
                continue
            if re.match(r'^(三|[四五六七八九十])', text) and ('判断题' in text or '简答题' in text or '论述题' in text):
                current_qtype = None
                i += 1
                continue

            # Parse numbered question
            if current_qtype:
                q_match = re.match(r'^(\d+)[\.\、\s\)]+(.+)', text)
                if q_match:
                    q_num = int(q_match.group(1))
                    q_text_part = q_match.group(2).strip()

                    # Options and answer might be inline or on following lines
                    # The format is mixed: sometimes options are inline on same line as Q,
                    # sometimes they're separate. Answers are 【答案】X

                    # Collect all text related to this question
                    q_full_text = q_text_part
                    options = []
                    answer_letters = []

                    # Check if answer is inline in the q_text_part
                    ans_in_text = re.search(r'【答案】\s*([A-G]+)', q_text_part)
                    if ans_in_text:
                        letters = ans_in_text.group(1)
                        answer_letters = [l for l in letters]
                        # Remove answer marker from question text
                        q_full_text = re.sub(r'【答案】.*$', '', q_text_part).strip()

                    j = i + 1
                    while j < len(paragraphs) and j < i + 15:
                        next_text = paragraphs[j]

                        # Check for section/break markers
                        if re.match(r'^(一、|二、|三、|四、|五、)', next_text) or \
                           any(next_text.startswith(jn) for jn in JIANG_NAMES):
                            break
                        if re.match(r'^\d+[\.\、\s\)]', next_text) and '【答案】' not in next_text:
                            break

                        # Check for answer marker
                        ans_match = re.search(r'【答案】\s*([A-G]+)', next_text)
                        if ans_match:
                            if not answer_letters:
                                answer_letters = [l for l in ans_match.group(1)]
                            # Also extract any option text before the answer marker
                            pre_ans = re.sub(r'【答案】.*$', '', next_text).strip()
                            if pre_ans:
                                opt_match = re.match(r'^([A-G])\s*[\.\、\）\)]?\s*(.+)', pre_ans)
                                if opt_match:
                                    options.append(opt_match.group(2).strip())
                                else:
                                    # This might be an inline option (A．xxx B．xxx C．xxx)
                                    inline_opts = re.findall(r'([A-G])\s*[\.\、\）\)]?\s*([^A-G]+?)(?=\s*[A-G]\s*[\.\、\）\)]|$)', pre_ans)
                                    for letter, content in inline_opts:
                                        content = content.strip().rstrip(';；')
                                        if content and letter not in [o[0] for o in inline_opts]:
                                            options.append(content)
                            j += 1
                            break

                        # Check for option
                        opt_match = re.match(r'^([A-G])\s*[\.\、\）\)]\s*(.+)', next_text)
                        if opt_match:
                            opt_content = opt_match.group(2).strip()
                            # Check if answer is inline
                            inline_ans = re.search(r'【答案】\s*([A-G]+)', opt_content)
                            if inline_ans:
                                opt_content = re.sub(r'【答案】.*$', '', opt_content).strip()
                                if not answer_letters:
                                    answer_letters = [l for l in inline_ans.group(1)]
                                options.append(opt_content)
                                j += 1
                                break
                            else:
                                options.append(opt_content)
                                j += 1
                                continue

                        # Could be option content as standalone text (PDF conversion artifact)
                        # e.g., "A．六千多万       B．七千多万       C．八千多万     D．九千多万"
                        inline_opts = re.findall(
                            r'([A-G])\s*[\.\、\）\)]?\s*(.+?)(?=\s+[A-G]\s*[\.\、\）\)]|\s*【答案】|$)',
                            next_text
                        )
                        if len(inline_opts) >= 2:
                            for letter, content in inline_opts:
                                content = content.strip().rstrip(';；')
                                ans_in = re.search(r'【答案】\s*([A-G]+)', content)
                                if ans_in:
                                    content = re.sub(r'【答案】.*$', '', content).strip()
                                    if not answer_letters:
                                        answer_letters = [l for l in ans_in.group(1)]
                                if content:
                                    options.append(content)
                            j += 1
                            continue

                        # Continuation of question or something else - skip
                        j += 1

                    if options and answer_letters:
                        # Validate answer indices
                        valid_indices = []
                        for a in answer_letters:
                            idx = option_to_index(a)
                            if 0 <= idx < len(options):
                                valid_indices.append(idx)

                        if valid_indices:
                            q_type = current_qtype
                            # Sometimes multi-choice with single answer gets mislabeled
                            if len(valid_indices) > 1:
                                q_type = "multi"

                            all_questions.append({
                                "type": q_type,
                                "question": q_full_text[:500],
                                "options": options,
                                "answer": valid_indices,
                                "source": "23题库PDF",
                                "priority": 2,
                                "_chapter": current_chapter,
                            })

                    i = j
                    continue

            i += 1

    # Attach chapter to questions
    result = []
    for q in all_questions:
        ch = q.pop("_chapter", "guide")
        result.append((ch, q))

    print(f"Parsing 23题库.pdf.docx: extracted {len(result)} questions")
    return result


# ─── Chapter assignment for Tiku1 questions ───────────────────

# Keywords for rough chapter matching
CHAPTER_KEYWORDS = {
    "guide": ["新时代", "马克思主义中国化", "新的飞跃", "两个结合", "百年未有之大变局",
              "社会主要矛盾", "四个意识", "两个维护", "两个确立", "四个自信",
              "十四个坚持", "十个明确", "历史方位", "中国特色社会主义进入新时代"],
    "ch01": ["总任务", "中国特色社会主义道路", "中华民族伟大复兴", "中国梦", "两步走"],
    "ch02": ["中国式现代化", "现代化强国", "共同富裕", "人类文明新形态",
             "中华民族伟大复兴", "社会主义现代化"],
    "ch03": ["党的全面领导", "党的领导", "党中央", "两个维护", "政治建设", "党的领导制度"],
    "ch04": ["以人民为中心", "人民至上", "人民立场", "群众路线", "人民主体"],
    "ch05": ["全面深化改革", "改革开放", "深化改革", "国家治理体系和治理能力现代化",
             "社会主义市场经济", "制度型开放", "对外开放"],
    "ch06": ["高质量发展", "新发展理念", "新发展阶段", "新发展格局", "供给侧结构性改革",
             "经济体系", "实体经济", "现代化经济体系"],
    "ch07": ["教育", "科技", "人才", "创新驱动", "科教兴国", "人才强国", "自主创新"],
    "ch08": ["全过程人民民主", "人民代表大会", "政治制度", "民主政治", "统一战线",
             "协商民主", "人民当家作主"],
    "ch09": ["依法治国", "法治", "宪法", "法律体系", "法治体系", "法治道路"],
    "ch10": ["文化强国", "文化自信", "意识形态", "社会主义核心价值观", "中华优秀传统文化",
             "文化建设", "精神文明"],
    "ch11": ["民生", "社会建设", "社会保障", "脱贫攻坚", "共同富裕", "就业",
             "收入分配", "健康中国"],
    "ch12": ["生态文明", "绿水青山", "美丽中国", "绿色发展", "碳达峰", "碳中和",
             "人与自然和谐共生"],
    "ch13": ["国家安全", "总体国家安全观", "安全", "安全格局", "平安中国"],
    "ch14": ["国防", "军队", "强军", "人民军队", "国防和军队现代化"],
    "ch15": ["一国两制", "祖国统一", "台湾", "香港", "澳门", "和平统一"],
    "ch16": ["外交", "人类命运共同体", "一带一路", "全球治理", "和平发展",
             "大国外交", "全球发展倡议"],
    "ch17": ["全面从严治党", "自我革命", "党的建设", "反腐败", "党的纪律", "正风肃纪"],
}


def assign_chapter_by_content(question, qtype):
    """Assign a chapter based on question content keywords."""
    text = question.lower()

    scores = {}
    for ch_id, keywords in CHAPTER_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw.lower() in text)
        scores[ch_id] = score

    if scores:
        best = max(scores, key=scores.get)
        if scores[best] > 0:
            return best

    return "guide"  # default


# ─── Merge and deduplicate ────────────────────────────────────

def merge_questions(existing_data, new_questions_list):
    """
    Merge new questions into existing chapters.
    new_questions_list: list of tuples (chapter_id, question_dict) or just question_dict
    """
    chapters_map = {}
    for ch in existing_data["chapters"]:
        chapters_map[ch["id"]] = ch

    # Track all existing question texts for dedup
    all_existing_texts = {}  # chapter_id -> list of (normalized_text, question_obj)
    for ch in existing_data["chapters"]:
        all_existing_texts[ch["id"]] = []
        for q in ch["questions"]:
            all_existing_texts[ch["id"]].append((normalize_question(q["question"]), q))

    stats = {"added": 0, "duplicates": 0, "errors": 0}

    # Get the starting IDs for each chapter
    next_ids = {}
    for ch in existing_data["chapters"]:
        if ch["questions"]:
            last_id = ch["questions"][-1]["id"]
            # Extract number: e.g., "guide_q0747" → 747
            m = re.search(r'_q(\d+)$', last_id)
            if m:
                next_ids[ch["id"]] = int(m.group(1)) + 1
            else:
                next_ids[ch["id"]] = 1
        else:
            next_ids[ch["id"]] = 1

    for item in new_questions_list:
        if isinstance(item, tuple):
            chapter_id, q = item
        else:
            chapter_id = assign_chapter_by_content(q.get("question", ""), q.get("type", "single"))

        if chapter_id not in chapters_map:
            # Try to find best matching chapter
            chapter_id = assign_chapter_by_content(q.get("question", ""), q.get("type", "single"))
            if chapter_id not in chapters_map:
                chapter_id = "guide"

        # Ensure chapter exists
        if chapter_id not in chapters_map:
            print(f"  WARN: Unknown chapter '{chapter_id}', defaulting to guide")
            chapter_id = "guide"
            if chapter_id not in chapters_map:
                stats["errors"] += 1
                continue

        q_text = q.get("question", "")
        q_normalized = normalize_question(q_text)

        # Check for duplicates within the same chapter first
        is_dup = False
        for norm_text, existing_q in all_existing_texts.get(chapter_id, []):
            sim = similarity(q_normalized, norm_text)
            if sim > 0.85:
                # Keep existing if source is "23版题库"
                if existing_q.get("source") == "23版题库":
                    is_dup = True
                    break
                # Otherwise, replace if new is better (23版题库 priority)
                elif q.get("source") == "23版题库" or q.get("source") == "23题库PDF":
                    # Remove old, add new
                    chapters_map[chapter_id]["questions"].remove(existing_q)
                    all_existing_texts[chapter_id].remove((norm_text, existing_q))
                    break
                else:
                    is_dup = True
                    break

        # Also check across all chapters
        if not is_dup:
            for ch_id, texts in all_existing_texts.items():
                for norm_text, existing_q in texts:
                    sim = similarity(q_normalized, norm_text)
                    if sim > 0.85:
                        if existing_q.get("source") == "23版题库":
                            is_dup = True
                            break
                        elif q.get("source") in ("23版题库", "23题库PDF"):
                            chapters_map[ch_id]["questions"].remove(existing_q)
                            texts.remove((norm_text, existing_q))
                            break
                        else:
                            is_dup = True
                            break
                if is_dup:
                    break

        if is_dup:
            stats["duplicates"] += 1
            continue

        # Assign ID
        qid = f"{chapter_id}_q{next_ids[chapter_id]:04d}"
        next_ids[chapter_id] += 1

        q["id"] = qid
        chapters_map[chapter_id]["questions"].append(q)
        all_existing_texts[chapter_id].append((q_normalized, q))
        stats["added"] += 1

    # Rebuild result
    result = {
        "meta": existing_data["meta"],
        "chapters": []
    }

    total = 0
    for ch_id in sorted(chapters_map.keys()):
        ch = chapters_map[ch_id]
        total += len(ch["questions"])
        result["chapters"].append(ch)

    result["meta"]["totalQuestions"] = total
    result["meta"]["version"] = "1.3"

    return result, stats


# ─── Main ──────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("Loading existing questions.json...")
    existing = load_existing()
    print(f"Loaded {existing['meta']['totalQuestions']} questions across {len(existing['chapters'])} chapters")

    all_new = []
    total_skipped_doc = 0

    # ── Source 1: 16讲 docx files ──
    print("\n" + "=" * 60)
    print("SOURCE 1: 16-lecture docx files")
    print("=" * 60)
    questions_16, skipped_docs = parse_all_16lectures()
    total_skipped_doc += len(skipped_docs)
    print(f"\n16-lecture total: {len(questions_16)} questions from docx files")
    if skipped_docs:
        print(f"Skipped .doc files ({len(skipped_docs)}): {', '.join(skipped_docs)}")
    all_new.extend(questions_16)

    # ── Source 2: 题库(1).docx ──
    print("\n" + "=" * 60)
    print("SOURCE 2: 习近平新时代中国特色社会主义思想概论题库(1).docx")
    print("=" * 60)
    questions_tiku1 = parse_tiku1()
    # Assign chapters by content
    q_with_chapters = []
    for q in questions_tiku1:
        ch = assign_chapter_by_content(q.get("question", ""), q.get("type", "single"))
        q_with_chapters.append((ch, q))
    print(f"题库(1) total: {len(q_with_chapters)} questions")
    all_new.extend(q_with_chapters)

    # ── Source 3: 23题库.pdf.docx ──
    print("\n" + "=" * 60)
    print("SOURCE 3: 习概23题库.pdf.docx")
    print("=" * 60)
    questions_23pdf = parse_23tiku_pdf()
    print(f"23题库PDF total: {len(questions_23pdf)} questions")
    all_new.extend(questions_23pdf)

    # ── Merge ──
    print("\n" + "=" * 60)
    print("MERGING AND DEDUPLICATING")
    print("=" * 60)
    print(f"Total new questions to check: {len(all_new)}")

    result, stats = merge_questions(existing, all_new)

    # ── Summary ──
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"From 16-lecture:       {len(questions_16)} questions")
    print(f"From 题库(1):          {len(questions_tiku1)} questions")
    print(f"From 23题库PDF:        {len(questions_23pdf)} questions")
    print(f"Skipped .doc files:    {total_skipped_doc}")
    print(f"Total new extracted:   {len(all_new)} questions")
    print(f"Added (after dedup):   {stats['added']} questions")
    print(f"Duplicates removed:    {stats['duplicates']} questions")
    print(f"Errors:                {stats['errors']}")
    print(f"Final total:           {result['meta']['totalQuestions']} questions")

    # Per-chapter breakdown
    print("\nPer chapter:")
    for ch in result["chapters"]:
        single = sum(1 for q in ch["questions"] if q["type"] == "single")
        multi = sum(1 for q in ch["questions"] if q["type"] == "multi")
        by_source = {}
        for q in ch["questions"]:
            src = q.get("source", "unknown")
            by_source[src] = by_source.get(src, 0) + 1
        src_str = ", ".join(f"{k}:{v}" for k, v in sorted(by_source.items()))
        print(f"  {ch['id']} ({ch['name']}): {len(ch['questions'])} ({single}s+{multi}m) [{src_str}]")

    # ── Write ──
    print(f"\nWriting to: {EXISTING_JSON}")
    with open(EXISTING_JSON, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print("Done!")


if __name__ == "__main__":
    main()
