#!/usr/bin/env python3
"""
从多个 docx 源文件中提取简答题/论述题（大题），输出为 JSON 题库文件。

源文件（按优先级排列）:
  1. 课后题详解.docx — 课后习题详细答案（主源）
  2. 习思想课后题.docx — 课后习题题目
  3. 习思想练习题分页 有简答.docx — 含简答题
  4. 2024习思想简单题回忆版.docx — 2024真题回忆
  5. 习思想 真题2022.docx — 2022真题
  6. 重点常考总结.docx — 重点常考总结
  7. 重点.docx — 重点
  8. 习近平新时代中国特色社会主义思想概论复习补充参考(1).docx — 复习补充

输出: ../data/essay_questions.json
"""

import os
import re
import json
import hashlib
from pathlib import Path
from docx import Document
from difflib import SequenceMatcher

# ============================================================
# 配置
# ============================================================

SOURCE_DIR = Path(r"C:\Users\99114\Desktop\习概")
OUTPUT_DIR = Path(r"C:\Users\99114\Desktop\习概复习\data")
OUTPUT_FILE = OUTPUT_DIR / "essay_questions.json"

SOURCE_FILES = [
    (SOURCE_DIR / "03_题库_辅助" / "课后题详解.docx", "课后题详解", 1),
    (SOURCE_DIR / "03_题库_辅助" / "习思想课后题.docx", "课后题", 2),
    (SOURCE_DIR / "03_题库_辅助" / "习思想练习题分页 有简答.docx", "练习题", 3),
    (SOURCE_DIR / "05_真题" / "2024习思想简单题回忆版.docx", "2024真题", 4),
    (SOURCE_DIR / "05_真题" / "习思想 真题2022.docx", "2022真题", 5),
    (SOURCE_DIR / "04_重点预测" / "重点常考总结.docx", "重点常考总结", 6),
    (SOURCE_DIR / "04_重点预测" / "重点.docx", "重点", 7),
    (SOURCE_DIR / "04_重点预测" / "习近平新时代中国特色社会主义思想概论复习补充参考(1).docx", "复习补充", 8),
]

# 章节映射
CHAPTER_MAP = {
    "guide": {"id": "guide", "name": "导论", "patterns": [
        "习近平新时代", "思想概论", "历史地位", "两个结合", "六个必须坚持",
        "时代背景", "历史方位", "创立", "思想及", "指导地位", "思想的历史",
        "思想是", "思想创立", "思想创立的", "科学内涵", "理论体系",
        "核心要义", "行动指南", "根本遵循", "思想回答", "重大时代课题",
        "历史性成就", "四个意识", "两个维护", "两个确立",
    ]},
    "ch01": {"id": "ch01", "name": "第一章 新时代坚持和发展中国特色社会主义", "patterns": [
        "新时代坚持和发展", "中国特色社会主义进入", "主要矛盾", "新的历史方位",
        "坚持和发展", "四个自信", "五位一体", "四个全面", "基本路线",
        "基本方略", "根本成就", "理论逻辑", "历史逻辑",
    ]},
    "ch02": {"id": "ch02", "name": "第二章 以中国式现代化全面推进中华民族伟大复兴", "patterns": [
        "中国式现代化", "中华民族伟大复兴", "中国梦", "两步走",
        "战略安排", "人类文明新形态", "全面建成小康", "现代化强",
        "中国式现代化是", "人类文明新形态", "现代化道路",
        "和平发展道路", "人口规模", "现代化", "物质文明和精神文明",
        "人与自然和谐", "全体人民共同富裕的现代化",
    ]},
    "ch03": {"id": "ch03", "name": "第三章 坚持党的全面领导", "patterns": [
        "坚持党的领导", "党的全面领导", "最本质特征", "两个维护",
        "四个意识", "党中央", "集中统一领导", "中国最大的国情",
        "党政军民学", "东西南北中", "党的核心", "党的领导是",
        "党是领导", "坚持和加强党的", "党的政治建设",
    ]},
    "ch04": {"id": "ch04", "name": "第四章 坚持以人民为中心", "patterns": [
        "以人民为中心", "共同富裕", "江山就是人民", "群众路线",
        "人民立场", "人民至上", "为人民服务", "以百姓心为心",
        "人民是", "人民的", "全体人民", "民生福祉", "人民的江山",
    ]},
    "ch05": {"id": "ch05", "name": "第五章 全面深化改革开放", "patterns": [
        "改革开放", "全面深化改革", "治理体系", "治理能力现代化",
        "制度优势", "开放", "深化改革", "中国开放的大门", "改革是",
        "体制机制", "制度建设",
    ]},
    "ch06": {"id": "ch06", "name": "第六章 推动高质量发展", "patterns": [
        "高质量发展", "新发展理念", "新发展格局", "供给侧",
        "经济", "市场", "分配制度", "现代化经济体系", "高质量发展是",
        "国内大循环", "双循环", "社会主义市场经济", "新发展阶段",
        "创新驱动", "区域协调",
    ]},
    "ch07": {"id": "ch07", "name": "第七章 社会主义现代化建设的教育科技人才战略", "patterns": [
        "科技", "教育强国", "人才强国", "科教兴国", "人才",
        "创新驱动发展", "科技自立自强", "高水平科技", "科技强国",
        "人才支撑", "教育科技人才",
    ]},
    "ch08": {"id": "ch08", "name": "第八章 发展全过程人民民主", "patterns": [
        "人民当家作主", "全过程人民民主", "民主", "政治制度",
        "人民代表大会", "协商民主", "统一战线", "政治发展道路",
        "社会主义民主", "最广泛最真实最管用", "人民民主",
    ]},
    "ch09": {"id": "ch09", "name": "第九章 全面依法治国", "patterns": [
        "依法治国", "法治", "宪法", "法律体系", "法治中国",
        "中国特色社会主义法治", "法治道路", "法治体系", "依宪治国",
        "法治国家", "全面依法", "科学立法", "严格执法", "公正司法",
    ]},
    "ch10": {"id": "ch10", "name": "第十章 建设社会主义文化强国", "patterns": [
        "文化自信", "意识形态", "核心价值观", "传统文化", "文化强国",
        "文化建设", "社会主义文化", "中华优秀传统", "文化繁荣",
        "精神文明建设", "文化软实力", "中华文化", "价值观念",
    ]},
    "ch11": {"id": "ch11", "name": "第十一章 以保障和改善民生为重点加强社会建设", "patterns": [
        "民生", "社会保障", "社会治理", "就业", "分配",
        "健康中国", "社会建设", "教育公平", "医疗", "住房",
        "脱贫攻坚", "社会公平", "公共服务", "基层治理",
    ]},
    "ch12": {"id": "ch12", "name": "第十二章 建设社会主义生态文明", "patterns": [
        "生态文明", "绿水青山", "绿色发展", "碳达峰", "碳中和",
        "环境", "生态保护", "美丽中国", "人与自然和谐",
        "生态环境", "绿色低碳", "污染防治",
    ]},
    "ch13": {"id": "ch13", "name": "第十三章 维护和塑造国家安全", "patterns": [
        "国家安全", "总体国家安全观", "安全", "风险挑战",
        "统筹发展和安全", "政治安全", "国家安全体系",
    ]},
    "ch14": {"id": "ch14", "name": "第十四章 建设巩固国防和强大人民军队", "patterns": [
        "国防", "军队", "强军", "军事", "人民军队",
        "党指挥枪", "能打胜仗", "作风优良", "强军目标",
    ]},
    "ch15": {"id": "ch15", "name": "第十五章 坚持\"一国两制\"和推进祖国完全统一", "patterns": [
        "一国两制", "祖国统一", "港澳", "台湾", "香港", "澳门",
        "祖国完全统一", "\"一国两制\"", "国家统一",
    ]},
    "ch16": {"id": "ch16", "name": "第十六章 中国特色大国外交和推动构建人类命运共同体", "patterns": [
        "外交", "人类命运共同体", "一带一路", "和平发展", "全球治理",
        "大国外交", "独立自主", "和平外交", "全球发展", "全球安全",
        "全球文明", "国际", "世界",
    ]},
    "ch17": {"id": "ch17", "name": "第十七章 全面从严治党", "patterns": [
        "全面从严治党", "党的建设", "反腐败", "自我革命",
        "党的纪律", "从严治党", "政治建设", "组织建设",
        "作风建设", "纪律建设", "制度治党", "党的自我革命",
        "跳出历史周期率", "反腐败斗争",
    ]},
}


def classify_chapter(text: str, filename: str = "") -> str:
    """根据文本内容判断所属章节"""
    text_clean = text.strip()
    scores = {}
    for ch_id, ch_info in CHAPTER_MAP.items():
        score = 0
        for pat in ch_info["patterns"]:
            if pat in text_clean:
                score += 1
        if score > 0:
            scores[ch_id] = score

    if not scores:
        return "guide"  # 默认归入导论

    # 返回得分最高的章节
    return max(scores, key=scores.get)


def normalize_text(text: str) -> str:
    """规范化文本：去除多余空白"""
    text = text.strip()
    text = re.sub(r'\s+', '', text)
    return text


def similarity(a: str, b: str) -> float:
    """计算两段文本的相似度"""
    return SequenceMatcher(None, normalize_text(a[:300]), normalize_text(b[:300])).ratio()


def extract_questions_from_keti_xiangjie(doc: Document, filename: str) -> list:
    """
    解析 课后题详解.docx
    格式: 第X章 -> 问答题 -> 编号题目 -> 【答案】... -> 下一题
    """
    questions = []
    current_chapter = None
    current_q_num = None
    current_q_text = []
    current_a_text = []
    in_answer = False
    in_question = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 检测章节标题
        chap_match = re.match(r'^第([一二三四五六七八九十\d]+)章\s+', text)
        if chap_match:
            # 保存前一道题
            if current_q_text and current_a_text:
                questions.append(_build_question(
                    current_q_text, current_a_text,
                    current_chapter, filename
                ))

            current_chapter = _parse_chapter_num(chap_match.group(1))
            current_q_text = []
            current_a_text = []
            in_answer = False
            in_question = False
            continue

        # 检测导论
        if '导论' in text[:5] or '前言' in text[:5]:
            if current_q_text and current_a_text:
                questions.append(_build_question(
                    current_q_text, current_a_text,
                    current_chapter, filename
                ))
            current_chapter = "guide"
            current_q_text = []
            current_a_text = []
            in_answer = False
            in_question = False
            continue

        # 跳过前言/目录内容
        if any(kw in text for kw in ['内容简介', '前言', '目录', '编者']):
            continue

        # 检测 "一、 问答题" / "二、 材料分析题" 等标题
        if re.match(r'^[一二三四五六七八九十]、\s*(问答题|简答题|论述题|材料分析)', text):
            in_question = True
            continue

        # 检测题目开始（数字编号）
        q_match = re.match(r'^(\d+)[\.\、\)]\s*(.+)', text)
        if q_match and len(q_match.group(2)) > 8 and not text.startswith('【答案】'):
            # 保存前一道题
            if current_q_text and current_a_text:
                questions.append(_build_question(
                    current_q_text, current_a_text,
                    current_chapter, filename
                ))
            current_q_num = q_match.group(1)
            current_q_text = [q_match.group(2)]
            current_a_text = []
            in_answer = False
            in_question = True
            continue

        # 检测答案开始
        if text.startswith('【答案】') or text.startswith('答：') or text.startswith('参考答案'):
            in_answer = True
            answer_start = text
            if text.startswith('【答案】'):
                answer_start = text[4:].strip()
            elif text.startswith('答：'):
                answer_start = text[2:].strip()
            else:
                answer_start = text[4:].strip()
            if answer_start:
                current_a_text.append(answer_start)
            continue

        # 检测下一节（如 "二、 材料分析题" 之后的题目，编号重置）
        if re.match(r'^[一二三四五六七八九十]、\s*', text):
            in_question = False
            in_answer = False
            continue

        # 积累文本
        if in_answer:
            current_a_text.append(text)
        elif in_question and current_q_text:
            # If we see what looks like a new question number mid-text
            if re.match(r'^\d+[\.\、\)]', text) and len(text) < 5:
                continue
            current_q_text.append(text)

    # 保存最后一道题
    if current_q_text and current_a_text:
        questions.append(_build_question(
            current_q_text, current_a_text,
            current_chapter, filename
        ))

    return questions


def extract_questions_from_keti_mulu(doc: Document, filename: str) -> list:
    """
    解析 习思想课后题.docx
    格式: 8.1问题... / 8.2问题...（简短题号+问题+简短答案混排）
    """
    questions = []
    current_q = []
    current_a = []
    in_answer = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检测 Q: 章节号.题号 开头
        q_match = re.match(r'^(\d+)\.(\d+)[\.\、\)]?\s*(.+)', text)
        if q_match:
            # 保存前一道题
            if current_q:
                questions.append(_build_question(
                    current_q, current_a,
                    _chapter_from_number(q_match.group(1)), filename
                ))
            current_q = [q_match.group(3)]
            current_a = []
            in_answer = False
            continue

        # 检测 答：开头
        if text.startswith('答：') or text.startswith('答:'):
            in_answer = True
            ans = text[2:].strip()
            if ans:
                current_a.append(ans)
            continue

        # 如果是新段落，且前一个 Q 在收集答案中
        if current_q:
            if in_answer:
                current_a.append(text)
            else:
                # 可能是继续的题目文本或答案
                current_q.append(text)

    if current_q:
        questions.append(_build_question(
            current_q, current_a,
            "guide", filename
        ))

    return questions


def extract_questions_from_lianxi(doc: Document, filename: str) -> list:
    """
    解析 习思想练习题分页 有简答.docx
    格式: 包含简答题/论述题区段，Q&A 以特定标记分隔
    """
    questions = []
    current_chapter = None
    current_q_text = []
    current_a_text = []
    in_answer = False
    in_question = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 章节检测
        chap_match = re.match(r'^第([一二三四五六七八九十\d]+)[章节]\s*', text)
        if chap_match:
            if current_q_text and current_a_text:
                questions.append(_build_question(
                    current_q_text, current_a_text,
                    current_chapter, filename
                ))
            current_chapter = _parse_chapter_num(chap_match.group(1))
            current_q_text = []
            current_a_text = []
            in_answer = False
            in_question = False
            continue

        # 检测 简答题/论述题 标题
        if any(kw in text for kw in ['简答题', '论述题', '问答题', '材料分析题']):
            in_question = True
            in_answer = False
            continue

        # 检测题目（数字开头）
        q_match = re.match(r'^(\d+)[\.\、\)]\s*(.+)', text)
        if q_match and not in_answer:
            if current_q_text and current_a_text:
                questions.append(_build_question(
                    current_q_text, current_a_text,
                    current_chapter, filename
                ))
            current_q_text = [q_match.group(2)]
            current_a_text = []
            in_answer = False
            in_question = True
            continue

        # 检测答案
        if re.match(r'^(答[：: ]|【答案】|参考答案)', text):
            in_answer = True
            ans = re.sub(r'^(答[：: ]|【答案】|参考答案)\s*', '', text)
            if ans:
                current_a_text.append(ans)
            continue

        # 检测类别标题结束
        if re.match(r'^[一二三四五六七八九十]、', text) and len(text) < 20:
            if in_answer and current_q_text and current_a_text:
                questions.append(_build_question(
                    current_q_text, current_a_text,
                    current_chapter, filename
                ))
                current_q_text = []
                current_a_text = []
            in_question = False
            in_answer = False
            continue

        # 积累文本
        if in_answer and current_q_text:
            current_a_text.append(text)
        elif in_question and current_q_text:
            current_q_text.append(text)

    if current_q_text and current_a_text:
        questions.append(_build_question(
            current_q_text, current_a_text,
            current_chapter, filename
        ))

    return questions


def extract_questions_from_2024_zhenti(doc: Document, filename: str) -> list:
    """
    解析 2024习思想简单题回忆版.docx
    格式: 题目列表（只有题目，没有答案）
    """
    questions = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 10:
            continue

        # 跳过文件路径和无关内容
        if any(kw in text[:20] for kw in ['文件', '编辑', '格式', '查看', '帮助', '微信', 'v:', 'V:']):
            continue

        # 看起来像是题目
        if '？' in text or '。' in text or '谈谈' in text or '如何' in text or '为什么' in text:
            # Clean the text
            text = re.sub(r'^\d+[\.\、\)]\s*', '', text)
            if len(text) > 10:
                questions.append(_build_question(
                    [text], [],
                    None, filename, qtype="简答", tags=["2024真题"]
                ))

    return questions


def extract_questions_from_2022_zhenti(doc: Document, filename: str) -> list:
    """
    解析 习思想 真题2022.docx
    格式: 试卷格式，包含选择题和可能的简答/论述
    """
    questions = []
    in_essay = False
    current_q = []
    current_a = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检测简答题/论述题标题
        if any(kw in text for kw in ['简答', '论述', '材料分析', '问答']):
            in_essay = True
            continue

        if not in_essay:
            continue

        # 检测题目
        q_match = re.match(r'^(\d+)[\.\、\)]\s*(.+)', text)
        if q_match and not text.startswith('答'):
            if current_q and current_a:
                questions.append(_build_question(
                    current_q, current_a,
                    None, filename, tags=["2022真题"]
                ))
            current_q = [q_match.group(2)]
            current_a = []
            continue

        if re.match(r'^(答[：:]|【答案】|参考答案)', text):
            ans = re.sub(r'^(答[：:]|【答案】|参考答案)\s*', '', text)
            if ans:
                current_a.append(ans)
            continue

        if current_q:
            if current_a:
                current_a.append(text)
            else:
                current_q.append(text)

    if current_q and current_a:
        questions.append(_build_question(
            current_q, current_a,
            None, filename, tags=["2022真题"]
        ))

    return questions


def extract_questions_from_zhongdian(doc: Document, filename: str) -> list:
    """
    解析 重点.docx 和 重点常考总结.docx
    格式: 章节标题 -> 知识点（含Q&A混排）
    特征: 题目常以 ? 结尾，答案紧跟其后
    """
    questions = []
    current_chapter = None
    current_q = []
    current_a = []
    in_answer = False
    current_tags = []
    current_type = "简答"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 跳过编辑/评论性文字
        if any(kw in text for kw in ["编者", "感谢", "仅供参考", "备考策略", "版权", "翻印"]):
            continue


        # 章节检测
        chap_match = re.match(r'^第([一二三四五六七八九十\d]+)[章节讲]\s*', text)
        if chap_match:
            if current_q and current_a:
                questions.append(_build_question(
                    current_q, current_a,
                    current_chapter, filename,
                    qtype=current_type, tags=current_tags + ["重点"]
                ))
            current_chapter = _parse_chapter_num(chap_match.group(1))
            current_q = []
            current_a = []
            in_answer = False
            current_tags = ["重点"]
            continue

        if '导论' in text[:5] or '前言' in text[:5]:
            if current_q and current_a:
                questions.append(_build_question(
                    current_q, current_a,
                    current_chapter, filename,
                    qtype=current_type, tags=current_tags + ["重点"]
                ))
            current_chapter = "guide"
            current_q = []
            current_a = []
            in_answer = False
            current_tags = ["重点"]
            continue

        # 检测标签：重点/了解/掌握/必考/常考
        tag_match = re.search(r'[（(](重点|必考|了解|掌握|常考|背诵)[）)]', text)
        if tag_match and len(text) < 50:
            tag = tag_match.group(1)
            if tag not in current_tags:
                current_tags.append(tag)
            if tag in ('重点', '必考', '常考', '背诵'):
                current_type = "论述"
            # If the text has content beyond the tag, treat the rest as question
            remaining = re.sub(r'[（(](重点|必考|了解|掌握|常考|背诵)[）)]', '', text).strip()
            if remaining and len(remaining) > 8:
                text = remaining  # process as question text below
            else:
                continue

        # 跳过纯标签行
        if re.match(r'^[（(](重点|必考|了解|掌握|常考|背诵|选择题|多选题)[）)]$', text):
            continue

        # 检测是否是新题目开始（以问号结尾的短句，或带有"一、"/"1、"等标记后问句）
        # 或者 "问：XXX" 格式
        if re.match(r'^问[：:]', text):
            # 新问题
            if current_q and current_a:
                questions.append(_build_question(
                    current_q, current_a,
                    current_chapter, filename,
                    qtype=current_type, tags=current_tags + ["重点"]
                ))
            current_q = [re.sub(r'^问[：:]\s*', '', text)]
            current_a = []
            in_answer = False
            continue

        # 以 "答：" 或 "答：" 开头
        if re.match(r'^答[：:]', text):
            in_answer = True
            ans = re.sub(r'^答[：:]?\s*', '', text)
            if ans:
                current_a.append(ans)
            continue

        # 检测编号开头的问题
        q_match = re.match(r'^(\d+)[\.\、\)]?\s*(.+)', text)
        if q_match and '？' in q_match.group(2):
            # 编号+问号结尾，判断为新问题
            if current_q and current_a:
                questions.append(_build_question(
                    current_q, current_a,
                    current_chapter, filename,
                    qtype=current_type, tags=current_tags + ["重点"]
                ))
            current_q = [q_match.group(2)]
            current_a = []
            in_answer = False
            continue

        # 如果当前在积累答案
        if in_answer and current_q:
            current_a.append(text)
            continue

        # 判断: 文本包含"？"且长度适中，可能是新题目
        if '？' in text and len(text) > 10 and len(text) < 200:
            if current_q and current_a:
                questions.append(_build_question(
                    current_q, current_a,
                    current_chapter, filename,
                    qtype=current_type, tags=current_tags + ["重点"]
                ))
            current_q = [text]
            current_a = []
            in_answer = False
            continue

        # 如果已经有题目在收集，可能是题目延续或答案
        if current_q:
            # 如果文本以数字列表或（1）、（2）等开头，可能是答案
            if re.match(r'^[（(]?\d+[）)]', text):
                in_answer = True
                current_a.append(text)
            elif in_answer:
                current_a.append(text)
            else:
                current_q.append(text)

    # 保存最后一道题
    if current_q and current_a:
        questions.append(_build_question(
            current_q, current_a,
            current_chapter, filename,
            qtype=current_type, tags=current_tags + ["重点"]
        ))

    return questions


def extract_questions_from_fuxi_buchong(doc: Document, filename: str) -> list:
    """
    解析 复习补充参考.docx
    格式: 提纲式复习资料，Q&A 混排
    """
    questions = []
    current_chapter = None
    current_q = []
    current_a = []
    in_answer = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 章节检测
        chap_match = re.match(r'^第([一二三四五六七八九十\d]+)[章节讲]\s*', text)
        if chap_match:
            if current_q and current_a:
                questions.append(_build_question(
                    current_q, current_a,
                    current_chapter, filename, tags=["复习补充"]
                ))
            current_chapter = _parse_chapter_num(chap_match.group(1))
            current_q = []
            current_a = []
            in_answer = False
            continue

        if '导论' in text[:5]:
            if current_q and current_a:
                questions.append(_build_question(
                    current_q, current_a,
                    current_chapter, filename, tags=["复习补充"]
                ))
            current_chapter = "guide"
            current_q = []
            current_a = []
            in_answer = False
            continue

        # 问题以 "l" 或 "-" 或数字标记开头（这是markdown风格的提纲）
        # 检查是否是问题描述
        if ('什么是' in text or '为什么' in text or '如何' in text or '怎样' in text or '？' in text) and len(text) > 15:
            if current_q and current_a:
                questions.append(_build_question(
                    current_q, current_a,
                    current_chapter, filename, tags=["复习补充"]
                ))
            # Clean the marker
            cleaned = re.sub(r'^[l\-\*\d\.、\s]+', '', text)
            current_q = [cleaned]
            current_a = []
            in_answer = False
            continue

        # 如果当前有题目，收集答案
        if current_q:
            if in_answer:
                current_a.append(text)
            elif len(text) > 10:
                in_answer = True
                current_a.append(text)
            else:
                in_answer = False
        else:
            # 可能是一个以问号结尾的短问题
            if '？' in text and len(text) > 10:
                current_q = [text]
                current_a = []
                in_answer = False

    if current_q and current_a:
        questions.append(_build_question(
            current_q, current_a,
            current_chapter, filename, tags=["复习补充"]
        ))

    return questions


def _build_question(q_parts: list, a_parts: list, chapter: str,
                    source: str, qtype: str = "简答",
                    tags: list = None) -> dict:
    """构建题目字典"""
    question = '\n'.join(q_parts).strip()
    answer = '\n'.join(a_parts).strip()

    # 清理题目文本
    question = re.sub(r'\s+', ' ', question).strip()
    answer = re.sub(r'\s+', ' ', answer).strip()

    # 截取
    question = question[:500]
    answer = answer[:3000]

    if not question or len(question) < 5:
        return None

    # 自动判定类型
    if not qtype or qtype == "简答":
        if any(kw in question for kw in ['论述', '分析', '谈谈', '理解', '阐述', '说明']):
            qtype = "论述"
        elif any(kw in question for kw in ['什么是', '定义', '含义', '概念', '名词']):
            qtype = "名词解释"
        elif len(question) > 100:
            qtype = "论述"

    # 自动判定章节
    if not chapter:
        chapter = classify_chapter(question, source)

    # 自动打标签
    final_tags = list(tags or [])
    if source == "2024真题":
        if "2024真题" not in final_tags:
            final_tags.append("2024真题")
    if source == "2022真题":
        if "2022真题" not in final_tags:
            final_tags.append("2022真题")
    if source == "课后题详解":
        if "课后题" not in final_tags:
            final_tags.append("课后题")
    if source in ("重点常考总结", "重点"):
        if "重点" not in final_tags:
            final_tags.append("重点")

    return {
        "question": question,
        "answer": answer,
        "chapter": chapter,
        "source": source,
        "type": qtype,
        "tags": final_tags,
    }


def _parse_chapter_num(ch_num_str: str) -> str:
    """将章节数字字符串转为 chapter ID"""
    mapping = {
        '一': '1', '二': '2', '三': '3', '四': '4', '五': '5',
        '六': '6', '七': '7', '八': '8', '九': '9', '十': '10',
        '十一': '11', '十二': '12', '十三': '13', '十四': '14',
        '十五': '15', '十六': '16', '十七': '17',
    }
    num = mapping.get(ch_num_str, ch_num_str)
    try:
        n = int(num)
        if 1 <= n <= 17:
            return f"ch{n:02d}"
    except ValueError:
        pass
    return "guide"


def _chapter_from_number(num_str: str) -> str:
    """从数字字符串推断章节"""
    try:
        n = int(num_str)
        if 8 <= n <= 17:
            return f"ch{n:02d}"
    except ValueError:
        pass
    return "guide"


def deduplicate_questions(questions: list, threshold: float = 0.72) -> list:
    """合并重复题目"""
    if not questions:
        return []

    # 按答案长度降序排列（保留答案更详细的）
    questions.sort(key=lambda q: len(q.get("answer", "")), reverse=True)

    unique = []
    for q in questions:
        is_dup = False
        for u in unique:
            sim = similarity(q["question"], u["question"])
            if sim >= threshold:
                is_dup = True
                # 合并标签
                for tag in q.get("tags", []):
                    if tag not in u.setdefault("tags", []):
                        u["tags"].append(tag)
                # 如果当前答案更长，替换
                if len(q.get("answer", "")) > len(u.get("answer", "")):
                    u["answer"] = q["answer"]
                    u["type"] = q.get("type", u.get("type"))
                break
        if not is_dup:
            unique.append(q)

    return unique


def assign_ids(questions: list) -> list:
    """为题目分配 ID 并按章节分组"""
    # 先按章节分组
    chapters = {}
    for q in questions:
        ch = q.get("chapter", "guide")
        if ch not in chapters:
            chapters[ch] = []
        chapters[ch].append(q)

    result = []
    for ch_id in sorted(chapters.keys()):
        ch_qs = chapters[ch_id]
        ch_info = CHAPTER_MAP.get(ch_id, {"id": ch_id, "name": ch_id})
        ch_entry = {
            "id": ch_info["id"],
            "name": ch_info["name"],
            "questions": []
        }

        for i, q in enumerate(ch_qs):
            qid = f"essay_{ch_id}_{i+1:03d}"
            entry = {
                "id": qid,
                "question": q["question"],
                "answer": q["answer"],
                "source": q.get("source", ""),
                "type": q.get("type", "简答"),
                "tags": q.get("tags", []),
            }
            ch_entry["questions"].append(entry)

        if ch_entry["questions"]:
            result.append(ch_entry)

    return result


def main():
    print("=" * 60)
    print("提取简答题/论述题（大题）")
    print("=" * 60)

    all_questions = []

    for filepath, filename, priority in SOURCE_FILES:
        if not filepath.exists():
            print(f"  [SKIP] 文件不存在: {filepath}")
            continue

        print(f"\n解析: {filepath.name}")

        try:
            doc = Document(str(filepath))
        except Exception as e:
            print(f"  [ERROR] 无法打开文件: {e}")
            continue

        para_count = len(doc.paragraphs)
        print(f"  段落数: {para_count}")

        questions = []

        if filename == "课后题详解":
            questions = extract_questions_from_keti_xiangjie(doc, filename)
        elif filename == "课后题":
            questions = extract_questions_from_keti_mulu(doc, filename)
        elif filename == "练习题":
            questions = extract_questions_from_lianxi(doc, filename)
        elif filename == "2024真题":
            questions = extract_questions_from_2024_zhenti(doc, filename)
        elif filename == "2022真题":
            questions = extract_questions_from_2022_zhenti(doc, filename)
        elif filename in ("重点常考总结", "重点"):
            questions = extract_questions_from_zhongdian(doc, filename)
        elif filename == "复习补充":
            questions = extract_questions_from_fuxi_buchong(doc, filename)
        else:
            # 通用解析器：尝试提取 Q&A
            questions = extract_questions_from_lianxi(doc, filename)

        # 过滤 None
        questions = [q for q in questions if q is not None and q.get("question")]
        print(f"  提取题目数: {len(questions)}")
        all_questions.extend(questions)

    print(f"\n总计提取: {len(all_questions)} 题")

    # 去重
    print("\n去重...")
    deduped = deduplicate_questions(all_questions)
    print(f"去重后: {len(deduped)} 题")

    # 分配 ID 并按章节组织
    print("\n分配ID并分组...")
    chapters = assign_ids(deduped)

    # 统计
    total_q = sum(len(ch["questions"]) for ch in chapters)
    print(f"最终: {len(chapters)} 个章节, {total_q} 道题")

    # 构建输出
    output = {
        "meta": {
            "subject": "习近平新时代中国特色社会主义思想概论",
            "shortName": "习概",
            "type": "essay",
            "version": "1.0",
        },
        "chapters": chapters,
    }

    # 确保输出目录存在
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # 写入 JSON
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"\n输出文件: {OUTPUT_FILE}")
    print(f"文件大小: {OUTPUT_FILE.stat().st_size / 1024:.1f} KB")

    # 打印各章节统计
    print("\n章节统计:")
    for ch in chapters:
        qc = len(ch["questions"])
        types = {}
        for q in ch["questions"]:
            t = q.get("type", "简答")
            types[t] = types.get(t, 0) + 1
        type_str = ", ".join(f"{k}:{v}" for k, v in types.items())
        print(f"  {ch['id']} {ch['name'][:20]}... — {qc}题 ({type_str})")


if __name__ == "__main__":
    main()
